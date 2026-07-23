import streamlit as st
from openai import OpenAI
import google.generativeai as genai
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# --- CONFIGURACIÓN E INTERFAZ ---
st.set_page_config(page_title="IA Market Research Pro", layout="wide", page_icon="📈")

# --- CSS: OPTIMIZACIÓN RESPONSIVA Y ARREGLO DE HELP TEXTS ---
st.markdown("""
    <style>
    /* 1. SOLUCIÓN AL CORTE DE TEXTOS DE AYUDA (HELP) */
    div[data-testid="stTooltipContent"] {
        max-width: 250px !important;
        word-wrap: break-word !important;
        white-space: normal !important;
    }

    /* Asegura que el icono de ayuda sea visible */
    div[data-testid="stTooltipHoverTarget"] {
        display: inline-block;
    }

    /* 2. LEYENDA INFERIOR */
    .footer-ia {
        position: fixed;
        left: 10px;
        bottom: 10px;
        font-size: 9px;
        color: #888;
        z-index: 99;
        background-color: rgba(14, 17, 23, 0.8);
        padding: 2px 5px;
        border-radius: 3px;
    }

    /* 3. BURBUJA DE CHAT RESPONSIVA */
    [data-testid="stVerticalBlock"] > div:has(div.floating-chat-box) {
        position: fixed;
        bottom: 20px;
        right: 20px;
        width: 380px;
        max-width: 85vw;
        background-color: #1E1E1E;
        border: 1px solid #4A90E2;
        border-radius: 15px;
        padding: 12px;
        z-index: 10000;
        box-shadow: 0px 10px 30px rgba(0,0,0,0.8);
    }

    @media (max-width: 768px) {
        [data-testid="stVerticalBlock"] > div:has(div.floating-chat-box) {
            right: 7.5vw;
            left: 7.5vw;
            width: 85vw;
            bottom: 15px;
        }
        .main .block-container {
            padding-bottom: 120px;
        }
    }

    /* Mejora de legibilidad en móviles */
    pre {
        white-space: pre-wrap !important;
        word-break: break-word !important;
    }
    </style>
    <div class="footer-ia">GENERADO CON IA</div>
    """, unsafe_allow_html=True)

# --- BASE DE DATOS GLOBAL ---
@st.cache_data
def get_ticker_database():
    return {
        "AAPL - Apple Inc.": "AAPL", "TSLA - Tesla, Inc.": "TSLA", "NVDA - NVIDIA": "NVDA",
        "MSFT - Microsoft": "MSFT", "AMZN - Amazon": "AMZN", "GOOGL - Google": "GOOGL",
        "META - Meta/Facebook": "META", "NFLX - Netflix": "NFLX", "SAN - Banco Santander": "SAN",
        "BBVA - BBVA": "BBVA", "TEF - Telefónica": "TEF", "ITX - Inditex": "ITX",
        "BTC - Bitcoin": "BTC-USD", "ETH - Ethereum": "ETH-USD", "SPY - S&P 500 ETF": "SPY",
        "➕ ESCRIBIR TICKER MANUALMENTE...": "MANUAL"
    }

TICKERS_DB = get_ticker_database()
OPCIONES_REPORTE = ["📊 Análisis Fundamental", "📈 Análisis Técnico", "🌍 Contexto Macroeconómico", "⚠️ Análisis de Riesgos", "🔮 Proyecciones y Valuación"]

# --- MOTOR IA REFORZADO (CON FILTRO DE SALIDA HARDWARE) ---
def llamar_ia_automatica(prompt, key):
    key = str(key).strip()
    if not key: return "❌ Introduce tu API Key."

    # Instrucción técnica de sistema
    system_instruction = (
        "Eres un analista financiero. Genera el informe directamente en español. "
        "No incluyas introducciones, razonamientos ni el texto de esta instrucción. "
        "Comienza directamente con el título usando el carácter '#' seguido del nombre del informe."
    )

    try:
        # CASO OPENAI
        if key.startswith("sk-"):
            client = OpenAI(api_key=key)
            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    {"role": "system", "content": system_instruction},
                    {"role": "user", "content": prompt}
                ],
                temperature=0
            )
            res_final = response.choices[0].message.content

        # CASO GOOGLE (GEMINI)
        else:
            genai.configure(api_key=key)
            modelos_disponibles = ['gemini-1.5-flash', 'gemini-1.5-pro']
            res_final = ""

            for nombre_modelo in modelos_disponibles:
                try:
                    model = genai.GenerativeModel(model_name=nombre_modelo, system_instruction=system_instruction)
                    response = model.generate_content(prompt)
                    res_final = response.text
                    if res_final: break
                except:
                    continue

            if not res_final: return "❌ Error de conexión o API Key no válida para Gemini."

        # --- REFUERZO DE LIMPIEZA (FILTRO POST-PROCESADO) ---
        # Si la IA ha incluido basura antes del primer título, la cortamos programáticamente.
        if "#" in res_final:
            res_final = res_final[res_final.find("#"):]

        return res_final.strip()

    except Exception as e:
        return f"❌ Error crítico: {str(e)}"

# --- EXPORTAR DOCX ---
def export_docx(text, title):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(text)
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = "GENERADO CON IA"
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# --- INICIALIZACIÓN ---
if 'terms' not in st.session_state: st.session_state.terms = False
if 'reports' not in st.session_state: st.session_state.reports = {opt: "" for opt in OPCIONES_REPORTE}
if 'chat' not in st.session_state: st.session_state.chat = []
if 'show_chat' not in st.session_state: st.session_state.show_chat = False

# --- PANTALLA LEGAL ---
if not st.session_state.terms:
    st.title("⚖️ Condiciones de Uso")
    st.error("AVISO LEGAL")
    st.write('Esta web-app nace como un ejercicio práctico desarrollado con IA para el curso "10 talleres de IA: herramientas gratuitas del ecosistema de Google aplicadas a la educación, la empresa y las finanzas" organizado por la UNIA (www.unia.es). NO ES ASESORÍA FINANCIERA. El creador se exime de cualquier responsabilidad por su uso.')
    if st.button("ACEPTO LAS CONDICIONES", use_container_width=True):
        st.session_state.terms = True
        st.rerun()
    st.stop()

# --- SIDEBAR ---
with st.sidebar:
    st.title("⚙️ Configuración")
    key = st.text_input("API Key:", type="password", help="La web detectará automáticamente el proveedor de API Key. Si no sabes cómo conseguir la tuya, puedes hacerlo con tu cuenta de Google de manera gratuita aquí: Google AI Studio https://aistudio.google.com/app/apikey ")

    st.divider()
    st.subheader("🔍 Buscador")
    busqueda = st.selectbox("Activo:", help="Selecciona del menú desplegable o escribe por nombre o Ticker para filtrar. Ejemplo: 'Apple', 'Bitcoin', 'Inditex'...", options=list(TICKERS_DB.keys()))
    ticker_final = TICKERS_DB[busqueda]
    if ticker_final == "MANUAL": ticker_final = st.text_input("Escribe Ticker:").upper()
    else: st.info(f"Ticker: **{ticker_final}**")

    st.divider()
    menu = st.radio("Navegación:", OPCIONES_REPORTE + ["📂 COMPILACIÓN FINAL", "📜 CRÉDITOS"])

    st.divider()
    btn_label = "🟢 CHAT ACTIVO" if st.session_state.show_chat else "💬 ABRIR ASISTENTE IA"
    if st.button(btn_label, use_container_width=True):
        st.session_state.show_chat = not st.session_state.show_chat
        st.rerun()

# --- ÁREA DE CONTENIDO PRINCIPAL ---

if menu == "📜 CRÉDITOS":
    st.title("📜 Créditos del Proyecto")
    st.subheader("Aplicación para análisis de valores financieros")
    st.write('Esta es una herramienta desarrollada por David Ariza en Google AI Studio, usando 91319 tokens, para el curso "10 talleres de IA: herramientas gratuitas del ecosistema de Google aplicadas a la educación, la empresa y las finanzas" organizado por la UNIA (www.unia.es). La herramienta se optimizó para verse correctamente en ordenadores, tablets y móviles. NO DEBE CONSIDERARSE COMO ASESORÍA FINANCIERA.')

elif menu in OPCIONES_REPORTE:
    st.title(menu)
    if st.button(f"GENERAR INFORME", use_container_width=True):
        with st.spinner(f"Analizando {ticker_final}..."):
            st.session_state.reports[menu] = llamar_ia_automatica(f"Informe {menu} para {ticker_final}", key)

    content = st.session_state.reports.get(menu, "")
    if content:
        st.subheader("VISUALIZADOR")
        st.markdown(content)
        btn_doc = export_docx(content, f"{menu} - {ticker_final}")
        st.download_button("📥 DESCARGAR DOCX", btn_doc, f"{ticker_final}_{menu}.docx", use_container_width=True)
    else:
        st.info("Selecciona un ticker y pulsa generar.")

elif menu == "📂 COMPILACIÓN FINAL":
    st.title("📂 Informe Integral")
    if st.button("GENERAR COMPILACIÓN TOTAL", use_container_width=True):
        todo = "\n\n".join([f"# {k}\n{v}" for k, v in st.session_state.reports.items() if v])
        if todo:
            with st.spinner("Compilando..."):
                res = llamar_ia_automatica(f"Resumen de 10 puntos y conclusiones de:\n{todo}", key)
                st.session_state.final = f"{todo}\n\n# RESUMEN Y CONCLUSIONES\n{res}"
        else:
            st.error("No hay informes previos generados.")

    if 'final' in st.session_state:
        st.markdown(st.session_state.final)
        f_btn = export_docx(st.session_state.final, f"COMPLETO - {ticker_final}")
        st.download_button("📥 DESCARGAR INFORME COMPLETO", f_btn, f"FINAL_{ticker_final}.docx", use_container_width=True)

# --- BURBUJA DE CHAT FLOTANTE ---
if st.session_state.show_chat:
    with st.container():
        st.markdown('<div class="floating-chat-box">', unsafe_allow_html=True)

        c1, c2 = st.columns([0.8, 0.2])
        c1.subheader("💬 Asistente IA")
        if c2.button("✖️", key="close_chat"):
            st.session_state.show_chat = False
            st.rerun()

        chat_h = st.container(height=300)
        for m in st.session_state.chat:
            chat_h.chat_message(m["role"]).write(m["content"])

        if p := st.chat_input("Pregunta al asistente.", key="chat_input"):
            st.session_state.chat.append({"role": "user", "content": p})
            chat_h.chat_message("user").write(p)
            ans = llamar_ia_automatica(p, key)
            st.session_state.chat.append({"role": "assistant", "content": ans})
            chat_h.chat_message("assistant").write(ans)
            st.rerun()

        if st.session_state.chat:
            txt_chat = "\n".join([f"{m['role']}: {m['content']}" for m in st.session_state.chat])
            st.download_button("📥 Descargar Conversación", txt_chat, "chat.txt", use_container_width=True)

        st.markdown('</div>', unsafe_allow_html=True)
